import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
import pytesseract
import os
import threading
import sys

base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))

# === CONFIGURACIÓN DEL SISTEMA ===
# Cambiar las siguientes rutas según la instalación local
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\Program Files\poppler-xx\Library\bin"


# === Textos en múltiples idiomas ===
texts = {
    "en": {
        "title": "OCR Bot 🧠",
        "file_label": "File (PDF or image):",
        "select_file": "Select file",
        "ocr_lang": "OCR Language:",
        "output_format": "Output format:",
        "output_name": "Output filename:",
        "save_folder": "Save folder:",
        "select_folder": "Select folder",
        "start_button": "Process OCR",
        "starting": "🚀 Starting...",
        "processing": "⏳ Processing: {}%",
        "processing_image": "⏳ Processing image... 100%",
        "done": "✅ Done!",
        "saved_as": "File saved as:\n{}\n\nOpen folder?",
        "error_title": "Error",
        "fill_all_fields": "Please fill in all fields.",
        "file_not_found": "File not found.",
        "language_menu": "Language:"
    },
    "es": {
        "title": "OCR Bot 🧠",
        "file_label": "Archivo (PDF o imagen):",
        "select_file": "Seleccionar archivo",
        "ocr_lang": "Idioma OCR:",
        "output_format": "Formato de salida:",
        "output_name": "Nombre del archivo de salida:",
        "save_folder": "Carpeta de guardado:",
        "select_folder": "Seleccionar carpeta",
        "start_button": "Procesar OCR",
        "starting": "🚀 Iniciando...",
        "processing": "⏳ Procesando: {}%",
        "processing_image": "⏳ Procesando imagen... 100%",
        "done": "✅ ¡Hecho!",
        "saved_as": "Archivo guardado en:\n{}\n\n¿Quieres abrir la carpeta?",
        "error_title": "Error",
        "fill_all_fields": "Completa todos los campos.",
        "file_not_found": "Archivo no encontrado.",
        "language_menu": "Idioma:"
    }
}

current_lang = "en"  # idioma por defecto

def update_language():
    global current_lang
    current_lang = language_var.get()
    lang = texts[current_lang]
    root.title(lang["title"])
    labels[0].config(text=lang["file_label"])
    file_button.config(text=lang["select_file"])
    labels[1].config(text=lang["ocr_lang"])
    labels[2].config(text=lang["output_format"])
    labels[3].config(text=lang["output_name"])
    labels[4].config(text=lang["save_folder"])
    folder_button.config(text=lang["select_folder"])
    start_button.config(text=lang["start_button"])
    labels[5].config(text=lang["language_menu"])

def ocr_image(image, lang_code):
    return pytesseract.image_to_string(image, lang=lang_code)

def save_as_txt(pages_text, output_path, output_name):
    full_path = os.path.join(output_path, f"{output_name}.txt")
    with open(full_path, "w", encoding="utf-8") as file:
        for i, text in enumerate(pages_text, 1):
            file.write(f"--- Page {i} ---\n{text}\n\n")
    return full_path

def save_as_docx(pages_text, output_path, output_name):
    full_path = os.path.join(output_path, f"{output_name}.docx")
    doc = Document()
    for i, text in enumerate(pages_text, 1):
        doc.add_heading(f"Page {i}", level=2)
        doc.add_paragraph(text)
    doc.save(full_path)
    return full_path

def process_ocr_thread(file_path, output_name, lang_code, format_choice, output_path):
    try:
        lang = texts[current_lang]
        pages_text = []

        if file_path.lower().endswith(".pdf"):
            pages = convert_from_path(file_path, poppler_path=POPPLER_PATH)
            total = len(pages)
            for i, page in enumerate(pages, 1):
                text = ocr_image(page, lang_code)
                pages_text.append(text)
                progress = int((i / total) * 100)
                progress_label.config(text=lang["processing"].format(progress))
                root.update_idletasks()
        else:
            img = Image.open(file_path)
            text = ocr_image(img, lang_code)
            pages_text.append(text)
            progress_label.config(text=lang["processing_image"])

        if format_choice == "txt":
            final_file = save_as_txt(pages_text, output_path, output_name)
        else:
            final_file = save_as_docx(pages_text, output_path, output_name)

        progress_label.config(text="")
        open_folder = messagebox.askyesno(lang["done"], lang["saved_as"].format(final_file))
        if open_folder:
            os.startfile(output_path)

    except Exception as e:
        messagebox.showerror(texts[current_lang]["error_title"], str(e))

def process_ocr():
    file_path = entry_file.get()
    output_name = entry_output.get()
    lang_code = lang_var.get()
    format_choice = format_var.get()
    output_path = output_dir.get()
    lang = texts[current_lang]

    if not file_path or not output_name or not output_path:
        messagebox.showerror(lang["error_title"], lang["fill_all_fields"])
        return

    if not os.path.exists(file_path):
        messagebox.showerror(lang["error_title"], lang["file_not_found"])
        return

    progress_label.config(text=lang["starting"])
    threading.Thread(target=process_ocr_thread, args=(file_path, output_name, lang_code, format_choice, output_path)).start()

# === INTERFAZ ===
root = tk.Tk()
root.geometry("450x420")
labels = []

labels.append(tk.Label(root))
labels[0].pack()
entry_file = tk.Entry(root, width=60)
entry_file.pack()
file_button = tk.Button(root, command=lambda: entry_file.insert(0, filedialog.askopenfilename()))
file_button.pack()

labels.append(tk.Label(root))
labels[1].pack()
lang_var = tk.StringVar(value="eng+spa")
tk.OptionMenu(root, lang_var, "eng", "spa", "eng+spa").pack()

labels.append(tk.Label(root))
labels[2].pack()
format_var = tk.StringVar(value="txt")
tk.OptionMenu(root, format_var, "txt", "docx").pack()

labels.append(tk.Label(root))
labels[3].pack()
entry_output = tk.Entry(root, width=40)
entry_output.pack()

labels.append(tk.Label(root))
labels[4].pack()
output_dir = tk.StringVar()
tk.Entry(root, textvariable=output_dir, width=60).pack()
folder_button = tk.Button(root, command=lambda: output_dir.set(filedialog.askdirectory()))
folder_button.pack()

progress_label = tk.Label(root, text="")
progress_label.pack(pady=10)

start_button = tk.Button(root, command=process_ocr, bg="green", fg="white", height=2)
start_button.pack(pady=10)

# Menú de idioma
labels.append(tk.Label(root))
labels[5].pack()
language_var = tk.StringVar(value="en")
tk.OptionMenu(root, language_var, "en", "es", command=lambda _: update_language()).pack()

# Aplicar idioma inicial
update_language()

root.mainloop()
